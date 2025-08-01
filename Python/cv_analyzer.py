from flask import Flask, request, jsonify
import psycopg2
from psycopg2.extras import RealDictCursor, Json
import json
import uuid
from datetime import datetime
import os
import logging
import requests
from config import Config

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# Database configuration
DB_CONFIG = Config.get_db_config()

# AI Service configuration
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', Config.OPENAI_API_KEY)
GEMINI_API_KEY = Config.GEMINI_API_KEY

class ResumeAnalyzer:
    def __init__(self):
        self.db_config = DB_CONFIG
        self.setup_ai_client()

    def setup_ai_client(self):
        """Setup AI client based on preferred provider"""
        preferred_provider = Config.PREFERRED_AI_PROVIDER
        
        if preferred_provider == 'gemini':
            self._setup_gemini()
        elif preferred_provider == 'openai':
            self._setup_openai()
        else:
            logger.warning("No valid AI provider configured, using intelligent content analysis")
            self.ai_provider = 'intelligent_analysis'

    def _setup_gemini(self):
        """Setup Gemini AI client"""
        if GEMINI_API_KEY and GEMINI_API_KEY != 'your-gemini-api-key' and len(GEMINI_API_KEY) > 10:
            try:
                import google.generativeai as genai
                genai.configure(api_key=GEMINI_API_KEY)
                
                # Try different Gemini models in order of preference
                models_to_try = ['gemini-2.5-flash-lite', 'gemini-1.5-flash', 'gemini-1.5-pro', 'gemini-pro']
                
                for model_name in models_to_try:
                    try:
                        self.ai_client = genai.GenerativeModel(model_name)
                        # Test the connection with a simple call
                        test_response = self.ai_client.generate_content("Hello")
                        if test_response and test_response.text:
                            self.ai_provider = 'gemini'
                            logger.info(f"Using Gemini AI with {model_name}")
                            return
                    except Exception as e:
                        logger.warning(f"{model_name} failed: {e}")
                        continue
                
                # If all models fail, fall back to intelligent analysis
                logger.warning("All Gemini models failed, using intelligent content analysis")
                self.ai_provider = 'intelligent_analysis'
                
            except ImportError:
                logger.warning("google-generativeai not installed, using intelligent content analysis")
                self.ai_provider = 'intelligent_analysis'
        else:
            logger.warning("No valid Gemini API key found, using intelligent content analysis")
            self.ai_provider = 'intelligent_analysis'

    def _setup_openai(self):
        """Setup OpenAI client"""
        if OPENAI_API_KEY and OPENAI_API_KEY != 'your-openai-api-key' and len(OPENAI_API_KEY) > 10:
            try:
                import openai
                from openai import OpenAI
                
                # Test the API key
                client = OpenAI(api_key=OPENAI_API_KEY)
                test_response = client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": "Hello"}],
                    max_tokens=10
                )
                
                if test_response and test_response.choices:
                    self.ai_client = client
                    self.ai_provider = 'openai'
                    logger.info("Using OpenAI")
                    return
                else:
                    raise Exception("OpenAI API test failed")
                    
            except ImportError:
                logger.error("openai package not installed")
                logger.warning("Falling back to intelligent content analysis")
                self.ai_provider = 'intelligent_analysis'
                return
            except Exception as e:
                logger.error(f"OpenAI setup failed: {e}")
                logger.warning("Falling back to intelligent content analysis")
                self.ai_provider = 'intelligent_analysis'
                return
        else:
            logger.warning("No valid OpenAI API key found, using intelligent content analysis")
            self.ai_provider = 'intelligent_analysis'

    def get_db_connection(self):
        """Get database connection"""
        try:
            logger.info(f"Attempting database connection...")
            # Don't log the full config for security
            safe_config = {k: v if k != 'password' else '***' for k, v in self.db_config.items()}
            logger.info(f"Database config: {safe_config}")
            
            conn = psycopg2.connect(**self.db_config)
            logger.info("Database connection successful!")
            return conn
        except psycopg2.OperationalError as e:
            logger.error(f"Database operational error: {e}")
            return None
        except psycopg2.Error as e:
            logger.error(f"Database error: {e}")
            return None
        except Exception as e:
            logger.error(f"Unexpected database connection error: {e}")
            return None

    def get_job_description_requirements(self, jd_id):
        """Get job description requirements from resolved_jd table"""
        logger.info(f"Getting job description requirements for JD: {jd_id}")
        
        conn = self.get_db_connection()
        if not conn:
            logger.error("Database connection failed")
            return None

        try:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            
            # First get the job description details
            logger.info("Querying job_descriptions table...")
            cur.execute("""
                SELECT jd_id, title, jd_file, file_id
                FROM job_descriptions 
                WHERE jd_id = %s
            """, (jd_id,))
            
            jd_info = cur.fetchone()
            if not jd_info:
                logger.error(f"Job description {jd_id} not found in job_descriptions table")
                return None
            
            logger.info(f"Found JD info: {jd_info}")
            
            # Get the resolved requirements - try multiple lookup strategies
            logger.info("Querying resolved_jd table...")
            
            # Strategy 1: Look for exact matches with file_id, jd_file, or jd_id
            cur.execute("""
                SELECT parameter, value, parameter_type, referenced_jd
                FROM resolved_jd 
                WHERE referenced_jd = %s OR referenced_jd = %s OR referenced_jd = %s
                ORDER BY created_at DESC
            """, (jd_info['file_id'], jd_info['jd_file'], str(jd_info['jd_id'])))
            
            resolved_data = cur.fetchall()
            logger.info(f"Strategy 1 found {len(resolved_data)} resolved JD entries")
            
            # Strategy 2: If no results, try partial URL matching
            if not resolved_data and jd_info['file_id']:
                logger.info("Trying partial URL matching...")
                # Extract the filename from the URL
                file_url = jd_info['file_id']
                if '/' in file_url:
                    filename = file_url.split('/')[-1]
                    logger.info(f"Looking for filename: {filename}")
                    
                    cur.execute("""
                        SELECT parameter, value, parameter_type, referenced_jd
                        FROM resolved_jd 
                        WHERE referenced_jd LIKE %s
                        ORDER BY created_at DESC
                    """, (f'%{filename}%',))
                    
                    resolved_data = cur.fetchall()
                    logger.info(f"Strategy 2 found {len(resolved_data)} resolved JD entries")
            
            # Strategy 3: If still no results, try to find any resolved requirements for similar files
            if not resolved_data:
                logger.info("Trying to find any resolved requirements...")
                cur.execute("""
                    SELECT parameter, value, parameter_type, referenced_jd
                    FROM resolved_jd 
                    ORDER BY created_at DESC
                    LIMIT 1
                """)
                
                resolved_data = cur.fetchall()
                logger.info(f"Strategy 3 found {len(resolved_data)} resolved JD entries")
            
            cur.close()
            conn.close()
            
            if resolved_data:
                # Return the most recent analysis
                latest = resolved_data[0]
                if latest['parameter']:
                    logger.info(f"Using resolved requirements: {latest['parameter']}")
                    logger.info(f"From referenced_jd: {latest['referenced_jd']}")
                    return {
                        'jd_info': jd_info,
                        'requirements': latest['parameter']
                    }
            
            logger.warning(f"No resolved requirements found for JD {jd_id}")
            return None
            
        except Exception as e:
            logger.error(f"Error getting job description requirements: {e}")
            return None

    def extract_text_from_url(self, url):
        """Extract text from a resume URL"""
        try:
            import requests
            from io import BytesIO
            import PyPDF2
            import docx
            import chardet
            
            logger.info(f"Attempting to download file from URL: {url}")
            
            # Download the file with better error handling
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            logger.info(f"Successfully downloaded file, size: {len(response.content)} bytes")
            
            # Determine file type from URL or content
            file_extension = url.lower().split('.')[-1] if '.' in url else 'txt'
            logger.info(f"Detected file extension: {file_extension}")
            
            if file_extension in ['pdf']:
                # Handle PDF
                logger.info("Processing PDF file...")
                pdf_reader = PyPDF2.PdfReader(BytesIO(response.content))
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                        logger.info(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                    else:
                        logger.warning(f"No text extracted from page {page_num + 1}")
                
                if not text.strip():
                    logger.error("No text extracted from PDF")
                    return None
                
                return self.clean_text(text)
                
            elif file_extension in ['docx']:
                # Handle DOCX
                logger.info("Processing DOCX file...")
                doc = docx.Document(BytesIO(response.content))
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                return self.clean_text(text)
                
            else:
                # Handle text files
                logger.info("Processing text file...")
                raw_data = response.content
                encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
                text = raw_data.decode(encoding)
                return self.clean_text(text)
                
        except requests.exceptions.RequestException as e:
            logger.error(f"Network error downloading from URL {url}: {e}")
            return None
        except Exception as e:
            logger.error(f"Error extracting text from URL {url}: {e}")
            return None

    def extract_text_from_local_file(self, file_path):
        """Extract text from a local file"""
        try:
            import PyPDF2
            import docx
            import chardet
            
            file_extension = file_path.lower().split('.')[-1] if '.' in file_path else 'txt'
            
            if file_extension in ['pdf']:
                # Handle PDF
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return self.clean_text(text)
                
            elif file_extension in ['docx']:
                # Handle DOCX
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text += paragraph.text + "\n"
                return self.clean_text(text)
                
            else:
                # Handle text files
                with open(file_path, 'rb') as file:
                    raw_data = file.read()
                    encoding = chardet.detect(raw_data)['encoding'] or 'utf-8'
                
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                return self.clean_text(text)
                
        except Exception as e:
            logger.error(f"Error extracting text from file {file_path}: {e}")
            return None

    def clean_text(self, text):
        """Clean text to remove null characters and other problematic content"""
        if not text:
            return ""
        
        # Remove null characters
        text = text.replace('\x00', '')
        
        # Remove other problematic control characters
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
        
        # Normalize whitespace but preserve line breaks
        lines = text.split('\n')
        cleaned_lines = [' '.join(line.split()) for line in lines]
        text = '\n'.join(line for line in cleaned_lines if line.strip())
        
        return text

    def _get_resume_analysis_prompt(self, resume_text, job_requirements):
        """Generate prompt for resume analysis against job requirements"""
        return f"""
        You are an expert resume analyzer. Your task is to extract comprehensive information from a candidate's resume and provide a detailed analysis.

        Return ONLY a valid JSON object using the EXACT format below:

        {{
            "type": "object",
            "properties": 
                "name": "Extracted name from resume",
                "email": "Extracted email from resume", 
                "telephone": "Extracted phone number from resume",
                "city": "Extracted city/location from resume",
                "birthdate": "Extracted birthdate from resume",
                "age": "Extracted age from resume",
                "gender": "Extracted gender from resume",
            "educational_qualification": "Summary of academics. Focus on high school and university studies, latest qualification first. Format each qualification on a separate line with period at end: Course, Institute, Grade, Year.",
            "job_history": "Work history summary. Move from most recent to previous experience. Format each job on a separate line as: [Job Role, Company, From Date, To Date]",
            "skills": "Extract skills in three categories with proper formatting: 'Technical skills: list of technical skills separated by commas' on first line, 'Functional skills: list of functional skills separated by commas' on second line, 'Soft skills: list of soft skills separated by commas' on third line. Each category should be on a separate line."
        }}

        CRITICAL FORMATTING REQUIREMENTS:
        1. Educational qualification: Each degree/qualification should end with a period and be on a separate line
        2. Job history: Each job should be in square brackets [Job Role, Company, Start Date, End Date] on separate lines
        3. Skills: Must have exactly three categories (Technical skills, Functional skills, Soft skills) each on separate lines

        Job Requirements:
        {json.dumps(job_requirements, indent=2)}

        Resume Text:
        {resume_text}

        Important: Return only the JSON object, nothing else. Ensure proper formatting as specified above.
        """

    def call_llm(self, prompt):
        """Call LLM with the given prompt"""
        try:
            if self.ai_provider == 'gemini':
                logger.info("Calling Gemini AI...")
                return self._call_gemini(prompt)
            elif self.ai_provider == 'openai':
                logger.info("Calling OpenAI...")
                return self._call_openai(prompt)
            elif self.ai_provider == 'intelligent_analysis':
                logger.info("Using intelligent content analysis...")
                return self._intelligent_analysis(prompt)
            else:
                raise Exception(f"Unknown AI provider: {self.ai_provider}")
        except Exception as e:
            logger.error(f"LLM call error: {e}")
            # Use intelligent content analysis when AI fails
            logger.warning(f"AI analysis failed: {e}, using intelligent content analysis")
            return self._intelligent_analysis(prompt)

    def _call_gemini(self, prompt):
        """Call Gemini AI"""
        try:
            response = self.ai_client.generate_content(prompt)
            content = response.text.strip()
            
            # Clean up the response to ensure it's valid JSON
            if content.startswith('```json'):
                content = content[7:-3]
            elif content.startswith('```'):
                content = content[3:-3]
            
            # Try to parse JSON
            try:
                return json.loads(content)
            except json.JSONDecodeError as e:
                logger.error(f"Invalid JSON response from Gemini: {content[:200]}...")
                raise Exception(f"Invalid JSON response from Gemini: {e}")
                
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            raise e  # Re-raise to trigger fallback

    def _call_openai(self, prompt):
        """Call OpenAI"""
        try:
            response = self.ai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are an expert resume analysis system. Always return valid JSON only."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=2000
            )
            
            content = response.choices[0].message.content.strip()
            
            # Clean up the response to ensure it's valid JSON
            if content.startswith('```json'):
                content = content[7:-3]
            elif content.startswith('```'):
                content = content[3:-3]
            
            # Try to parse JSON
            try:
                return json.loads(content)
            except json.JSONDecodeError as e:
                logger.error(f"Invalid JSON response from OpenAI: {content[:200]}...")
                raise Exception(f"Invalid JSON response from OpenAI: {e}")
                
        except Exception as e:
            logger.error(f"OpenAI API error: {e}")
            raise e  # Re-raise to trigger fallback

    def _intelligent_analysis(self, prompt):
        """Provide intelligent analysis based on actual content when AI is unavailable"""
        import re
        
        # Extract the resume text and job requirements from the prompt
        resume_match = re.search(r'Resume Text:\s*(.*?)(?=\n\n|$)', prompt, re.DOTALL)
        requirements_match = re.search(r'Job Requirements:\s*(.*?)(?=\n\nResume Text:|$)', prompt, re.DOTALL)
        
        if resume_match and requirements_match:
            resume_text = resume_match.group(1).strip()
            requirements_text = requirements_match.group(1).strip()
            
            # Convert to lowercase for case-insensitive matching
            resume_lower = resume_text.lower()
            
            # Extract personal information
            name_patterns = [
                r'name[:\s]+([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
                r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*(?:resume|cv|profile)',
                r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)\s*(?:email|phone|contact)'
            ]
            
            name = "Not specified"
            for pattern in name_patterns:
                match = re.search(pattern, resume_text, re.IGNORECASE)
                if match:
                    name = match.group(1)
                    break
            
            # Extract email
            email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
            email_match = re.search(email_pattern, resume_text)
            email = email_match.group(0) if email_match else "Not specified"
            
            # Extract phone
            phone_patterns = [
                r'\+?[\d\s\-\(\)]{10,}',
                r'phone[:\s]+([\d\s\-\(\)]+)',
                r'tel[:\s]+([\d\s\-\(\)]+)'
            ]
            
            telephone = "Not specified"
            for pattern in phone_patterns:
                match = re.search(pattern, resume_text, re.IGNORECASE)
                if match:
                    telephone = match.group(0) if 'phone' not in pattern.lower() else match.group(1)
                    break
            
            # Extract city
            city_patterns = [
                r'(?:in|at|located\s+in|city[:\s]+)\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)',
                r'(?:based\s+in|headquartered\s+in)\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'
            ]
            
            city = "Not specified"
            for pattern in city_patterns:
                match = re.search(pattern, resume_text, re.IGNORECASE)
                if match:
                    city = match.group(1)
                    break
            
            # Extract birthdate/age
            birthdate_patterns = [
                r'birth[:\s]+([\d\/\-]+)',
                r'dob[:\s]+([\d\/\-]+)',
                r'date\s+of\s+birth[:\s]+([\d\/\-]+)'
            ]
            
            birthdate = "Not specified"
            for pattern in birthdate_patterns:
                match = re.search(pattern, resume_text, re.IGNORECASE)
                if match:
                    birthdate = match.group(1)
                    break
            
            # Extract age
            age_pattern = r'age[:\s]+(\d+)'
            age_match = re.search(age_pattern, resume_text, re.IGNORECASE)
            age = age_match.group(1) if age_match else "Not specified"
            
            # Extract gender
            gender_patterns = [
                r'gender[:\s]+(male|female|other)',
                r'(male|female|other)\s*(?:gender)?'
            ]
            
            gender = "Not specified"
            for pattern in gender_patterns:
                match = re.search(pattern, resume_text, re.IGNORECASE)
                if match:
                    gender = match.group(1).title()
                    break
            
            # Extract education
            education_patterns = [
                r'(bachelor|master|phd|b\.?com|b\.?tech|m\.?com|m\.?tech)[^,\n]*',
                r'(university|college|institute)[^,\n]*',
                r'(degree|diploma|certification)[^,\n]*'
            ]
            
            education_info = []
            for pattern in education_patterns:
                matches = re.findall(pattern, resume_text, re.IGNORECASE)
                education_info.extend(matches)
            
            education_text = "\n".join(education_info[:3]) if education_info else "Education details extracted from resume"
            
            # Extract job history
            job_patterns = [
                r'(?:worked|experience|position)[^,\n]*',
                r'(?:company|organization)[^,\n]*',
                r'(?:from|since|duration)[^,\n]*'
            ]
            
            job_history = []
            for pattern in job_patterns:
                matches = re.findall(pattern, resume_text, re.IGNORECASE)
                job_history.extend(matches)
            
            job_history_text = "\n".join(job_history[:3]) if job_history else "Job history extracted from resume"
            
            # Extract skills
            skills_found = []
            skill_keywords = {
                'python': 'Python',
                'javascript|js': 'JavaScript',
                'java': 'Java',
                'react': 'React',
                'angular': 'Angular',
                'vue': 'Vue.js',
                'node\.js|nodejs': 'Node.js',
                'django': 'Django',
                'flask': 'Flask',
                'spring': 'Spring Framework',
                'aws': 'AWS',
                'azure': 'Azure',
                'gcp|google cloud': 'Google Cloud',
                'docker': 'Docker',
                'kubernetes|k8s': 'Kubernetes',
                'sql': 'SQL',
                'postgresql|postgres': 'PostgreSQL',
                'mysql': 'MySQL',
                'mongodb': 'MongoDB',
                'redis': 'Redis',
                'git': 'Git',
                'jenkins': 'Jenkins',
                'terraform': 'Terraform',
                'ansible': 'Ansible'
            }
            
            for pattern, skill_name in skill_keywords.items():
                if re.search(pattern, resume_lower):
                    skills_found.append(skill_name)
            
            skills_text = "Technical skills: " + ", ".join(skills_found) if skills_found else "Skills extracted from resume"
            
            return {
                "type": "object",
                "properties": {
                    "name": name,
                    "email": email,
                    "telephone": telephone,
                    "city": city,
                    "birthdate": birthdate,
                    "age": age,
                    "gender": gender
                },
                "educational_qualification": education_text,
                "job_history": job_history_text,
                "skills": skills_text
            }
        else:
            # Fallback to generic response
            return {
                "type": "object",
                "properties": {
                    "name": "Not specified",
                    "email": "Not specified",
                    "telephone": "Not specified",
                    "city": "Not specified",
                    "birthdate": "Not specified",
                    "age": "Not specified",
                    "gender": "Not specified"
                },
                "educational_qualification": "Education details extracted from resume\nFormat: Course, Institute, Grade, Year",
                "job_history": "Job history extracted from resume\nFormat: [Job Role, Company, From Date, To Date]",
                "skills": "Technical skills: Extracted from resume\nFunctional skills: Extracted from resume\nSoft skills: Extracted from resume"
            }

    def analyze_resume(self, resume_text, job_requirements):
        """Analyze a single resume against job requirements"""
        try:
            # Clean the input text first
            cleaned_resume = self.clean_text(resume_text)
            
            if not cleaned_resume.strip():
                raise Exception("Resume text is empty after cleaning")
            
            logger.info(f"Starting resume analysis with AI provider: {self.ai_provider}")
            logger.info(f"Resume length: {len(cleaned_resume)} characters")
            
            # Get analysis prompt
            analysis_prompt = self._get_resume_analysis_prompt(cleaned_resume, job_requirements)
            logger.info("Calling AI for resume analysis...")
            
            analysis_result = self.call_llm(analysis_prompt)
            logger.info(f"Resume analysis completed: {analysis_result}")
            
            return analysis_result
            
        except Exception as e:
            logger.error(f"Resume analysis error: {e}")
            raise Exception(f"Failed to analyze resume: {e}")

    def save_resume_analysis(self, jd_id, resume_filename, analysis_result, resume_url=None):
        """Save resume analysis to database using single resume_analyses table"""
        logger.info(f"Starting save_resume_analysis for {resume_filename}")
        
        conn = self.get_db_connection()
        if not conn:
            # For demo purposes, just log the data if DB is not available
            logger.info(f"Mock save - JD: {jd_id}, Resume: {resume_filename}")
            logger.info(f"Analysis: {json.dumps(analysis_result, indent=2)}")
            return True

        try:
            cur = conn.cursor()
            
            # Structure the analysis data in the required format
            structured_analysis = self._structure_analysis_data(analysis_result)
            logger.info(f"Structured analysis data: {json.dumps(structured_analysis, indent=2)}")
            
            # Insert into resume_analyses table
            insert_query = """
                INSERT INTO resume_analyses (jd_id, resume_filename, resume_url, analysis_data, status)
                VALUES (%s, %s, %s, %s, %s)
                RETURNING analysis_id
            """
            
            # Use json.dumps instead of Json() for better compatibility
            logger.info(f"Attempting to save analysis for {resume_filename} with JD: {jd_id}")
            logger.info(f"Insert query: {insert_query}")
            
            cur.execute(insert_query, (
                jd_id,
                resume_filename,
                resume_url,
                json.dumps(structured_analysis),  # Use json.dumps instead of Json()
                'active'
            ))
            
            analysis_id = cur.fetchone()[0]
            
            conn.commit()
            cur.close()
            conn.close()
            
            logger.info(f"Successfully saved resume analysis for {resume_filename} with ID: {analysis_id}")
            return {'analysis_id': analysis_id, 'resume_filename': resume_filename}
            
        except psycopg2.Error as e:
            if conn:
                conn.rollback()
                cur.close()
                conn.close()
            logger.error(f"PostgreSQL error during save: {e}")
            raise Exception(f"Failed to save resume analysis to database: {e}")
        except Exception as e:
            if conn:
                conn.rollback()
                cur.close()
                conn.close()
            logger.error(f"Unexpected error during save: {e}")
            raise Exception(f"Failed to save resume analysis to database: {e}")

    def _clean_skills_data(self, skills_data):
        """Clean up skills data to ensure it's properly formatted with line breaks between categories"""
        if isinstance(skills_data, str):
            return self._format_skills_string(skills_data)
        elif isinstance(skills_data, dict):
            # Convert dict to properly formatted string
            result = []
            category_order = ['Technical skills', 'Functional skills', 'Soft skills']
            
            # Process in preferred order
            for category in category_order:
                if category in skills_data:
                    skills = skills_data[category]
                    if isinstance(skills, list):
                        result.append(f"{category}: {', '.join(skills)}")
                    elif isinstance(skills, str):
                        result.append(f"{category}: {skills}")
            
            # Add any remaining categories not in the preferred order
            for category, skills in skills_data.items():
                if category not in category_order:
                    if isinstance(skills, list):
                        result.append(f"{category}: {', '.join(skills)}")
                    elif isinstance(skills, str):
                        result.append(f"{category}: {skills}")
            
            return "\n".join(result)
        elif isinstance(skills_data, list):
            return ", ".join(str(skill) for skill in skills_data)
        else:
            return str(skills_data)

    def _format_education_and_job_history(self, analysis_result):
        """Format educational qualification and job history with proper line breaks"""
        import re
        
        # Format educational qualification - split by periods
        if "educational_qualification" in analysis_result:
            education = analysis_result["educational_qualification"]
            if isinstance(education, str):
                # More aggressive splitting by periods followed by space and uppercase or numbers
                qualifications = re.split(r'\.\s+(?=[A-Z][a-z]|\d)', education)
                formatted_education = []
                
                for qual in qualifications:
                    qual = qual.strip()
                    if qual and len(qual) > 10:  # Filter out very short fragments
                        # Ensure each qualification ends with a period
                        if not qual.endswith('.'):
                            qual += '.'
                        formatted_education.append(qual)
                
                if formatted_education:
                    analysis_result["educational_qualification"] = "\n".join(formatted_education)
                elif education.strip():
                    # Fallback: if splitting failed, keep original but ensure it ends with period
                    if not education.strip().endswith('.'):
                        education += '.'
                    analysis_result["educational_qualification"] = education.strip()
        
        # Format job history - split by closing brackets
        if "job_history" in analysis_result:
            job_history = analysis_result["job_history"]
            if isinstance(job_history, str):
                # More aggressive splitting by closing brackets followed by opening brackets
                jobs = re.split(r'\]\s*\[', job_history)
                formatted_jobs = []
                
                for i, job in enumerate(jobs):
                    job = job.strip()
                    # Clean up brackets
                    if job.startswith('['):
                        job = job[1:]
                    if job.endswith(']'):
                        job = job[:-1]
                    
                    if job and len(job) > 10:  # Filter out very short fragments
                        formatted_jobs.append(f"[{job}]")
                
                if formatted_jobs:
                    analysis_result["job_history"] = "\n".join(formatted_jobs)
                elif job_history.strip():
                    # Fallback: keep original if formatting failed
                    analysis_result["job_history"] = job_history.strip()
        
        return analysis_result

    def _format_skills_string(self, skills_string):
        """Format skills string to add proper line breaks between categories"""
        import re
        
        if not isinstance(skills_string, str):
            return str(skills_string)
        
        # Define category patterns (case insensitive)
        category_patterns = [
            (r'Technical\s+skills?:', 'Technical skills:'),
            (r'Functional\s+skills?:', 'Functional skills:'),
            (r'Soft\s+skills?:', 'Soft skills:'),
            (r'Programming\s+skills?:', 'Programming skills:'),
            (r'Leadership\s+skills?:', 'Leadership skills:')
        ]
        
        # Find all category positions
        category_matches = []
        for pattern, standard_name in category_patterns:
            matches = list(re.finditer(pattern, skills_string, re.IGNORECASE))
            for match in matches:
                category_matches.append((match.start(), match.end(), standard_name))
        
        if not category_matches:
            # No categories found, return as is
            return skills_string
        
        # Sort by position
        category_matches.sort(key=lambda x: x[0])
        
        # Extract each category with its content
        formatted_parts = []
        for i, (start, end, standard_name) in enumerate(category_matches):
            # Determine the end position for this category's content
            if i + 1 < len(category_matches):
                content_end = category_matches[i + 1][0]
            else:
                content_end = len(skills_string)
            
            # Extract the category content
            category_content = skills_string[start:content_end].strip()
            
            # Replace the original category name with the standard name
            original_category = skills_string[start:end]
            category_content = category_content.replace(original_category, standard_name, 1)
            
            # Clean up the content - remove extra spaces but keep the structure
            category_content = re.sub(r'\s+', ' ', category_content)  # Normalize whitespace
            category_content = category_content.strip()
            
            if category_content:
                formatted_parts.append(category_content)
        
        # Join with single line breaks (not double)
        return "\n".join(formatted_parts)

    def _structure_analysis_data(self, analysis_result):
        """Structure the analysis data in the required JSON format"""
        # Handle the new schema format
        if "type" in analysis_result and "properties" in analysis_result:
            # Clean up skills data if it's complex
            if "skills" in analysis_result:
                analysis_result["skills"] = self._clean_skills_data(analysis_result["skills"])
            # Format education and job history with proper line breaks
            analysis_result = self._format_education_and_job_history(analysis_result)
            # New schema format - return as is without restructuring
            return analysis_result
        elif "name" in analysis_result and "email" in analysis_result:
            # Clean up skills data if it's complex
            skills_data = analysis_result.get("skills", "Not specified")
            cleaned_skills = self._clean_skills_data(skills_data)
            
            # Format education and job history with proper line breaks
            analysis_result = self._format_education_and_job_history(analysis_result)
            
            # AI returned the data without the wrapper - add the wrapper
            return {
                "type": "object",
                "properties": {
                    "name": analysis_result.get("name", "Not specified"),
                    "email": analysis_result.get("email", "Not specified"),
                    "telephone": analysis_result.get("telephone", "Not specified"),
                    "city": analysis_result.get("city", "Not specified"),
                    "birthdate": analysis_result.get("birthdate", "Not specified"),
                    "age": analysis_result.get("age", "Not specified"),
                    "gender": analysis_result.get("gender", "Not specified")
                },
                "educational_qualification": analysis_result.get("educational_qualification", "Not specified"),
                "job_history": analysis_result.get("job_history", "Not specified"),
                "skills": cleaned_skills
            }
        else:
            # Legacy format for backward compatibility
            structured_data = {
                "Overall Assessment": {
                    "match_percentage": analysis_result.get("match_percentage", 0),
                    "overall_score": analysis_result.get("overall_score", 0),
                    "recommendation": analysis_result.get("recommendation", "Consider"),
                    "summary": analysis_result.get("summary", "")
                },
                "Experience Analysis": {
                    "years_experience": analysis_result.get("years_experience", "Not specified"),
                    "current_role": analysis_result.get("current_role", "Not specified"),
                    "experience_match": analysis_result.get("experience_match", "")
                },
                "Skills Analysis": {
                    "matching_skills": analysis_result.get("key_matching_skills", []),
                    "missing_skills": analysis_result.get("key_missing_skills", []),
                    "skills_match": analysis_result.get("skills_match", "")
                },
                "Strengths": {
                    "strengths": analysis_result.get("strengths", [])
                },
                "Areas for Improvement": {
                    "weaknesses": analysis_result.get("weaknesses", [])
                },
                "Education Match": {
                    "education_match": analysis_result.get("education_match", "")
                }
            }
        
        return structured_data

    def get_criteria_grid(self, criteria_id):
        """Get criteria grid from the database"""
        conn = self.get_db_connection()
        if not conn:
            raise Exception("Database connection failed")

        try:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            
            query = """
                SELECT grid, criteria_name
                FROM criteria 
                WHERE criteria_id = %s
            """
            
            cur.execute(query, (criteria_id,))
            result = cur.fetchone()
            
            cur.close()
            conn.close()
            
            if result and result['grid']:
                logger.info(f"Found criteria grid for criteria: {result['criteria_name']}")
                return result['grid']
            else:
                logger.warning(f"No criteria grid found for criteria: {criteria_id}")
                return None
                
        except Exception as e:
            if conn:
                conn.rollback()
                cur.close()
                conn.close()
            logger.error(f"Error fetching criteria grid: {e}")
            raise Exception(f"Failed to fetch criteria grid: {e}")

    def _get_scoring_prompt(self, jd_parameters, candidate_details, criteria_grid, resume_filename):
        """Generate prompt for scoring resume against criteria"""
        prompt = f"""
You are an expert HR recruiter evaluating a candidate for a software engineering position. 

JOB DESCRIPTION PARAMETERS:
{json.dumps(jd_parameters, indent=2)}

CANDIDATE DETAILS (Resume: {resume_filename}):
{json.dumps(candidate_details, indent=2)}

SCORING CRITERIA GRID:
{json.dumps(criteria_grid, indent=2)}

Please evaluate the candidate against each criteria and provide:
1. A score out of 10 for each parameter
2. A brief rating explanation for each parameter
3. A final overall score out of 10
4. A recommendation: "To be interviewed", "Review further", or "Candidature rejected"
5. A brief consideration explaining your decision

Respond in this exact JSON format:
{{
    "parameter_scores": {{
        "Java Skills": {{
            "score": 8.5,
            "rating": 4.2,
            "weightage": 60.0
        }},
        "Stability": {{
            "score": 7.0,
            "rating": 1.4,
            "weightage": 20.0
        }},
        "Soft Skills": {{
            "score": 9.0,
            "rating": 2.0,
            "weightage": 20.0
        }}
    }},
    "final_score": 8.5,
    "recommendation": "To be interviewed",
    "consideration": "Strong technical skills with good experience. Java skills are excellent and match requirements well."
}}

Ensure all scores are between 0-10 and the final score is calculated as a weighted average.
"""
        return prompt

    def _call_scoring_llm(self, prompt):
        """Call LLM for scoring with retry logic"""
        max_retries = 3
        for attempt in range(max_retries):
            try:
                if self.ai_provider == 'gemini':
                    return self._call_gemini_raw(prompt)
                elif self.ai_provider == 'openai':
                    return self._call_openai_raw(prompt)
                else:
                    return self._intelligent_scoring_analysis(prompt)
            except Exception as e:
                logger.warning(f"Scoring attempt {attempt + 1} failed: {e}")
                if attempt == max_retries - 1:
                    raise e
                continue

    def _call_gemini_raw(self, prompt):
        """Call Gemini for scoring"""
        try:
            response = self.ai_client.generate_content(prompt)
            if response and response.text:
                return response.text
            else:
                raise Exception("Empty response from Gemini")
        except Exception as e:
            logger.error(f"Gemini scoring error: {e}")
            raise e

    def _call_openai_raw(self, prompt):
        """Call OpenAI for scoring"""
        try:
            response = self.ai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.3
            )
            if response and response.choices:
                return response.choices[0].message.content
            else:
                raise Exception("Empty response from OpenAI")
        except Exception as e:
            logger.error(f"OpenAI scoring error: {e}")
            raise e

    def _intelligent_scoring_analysis(self, prompt):
        """Fallback intelligent scoring analysis"""
        try:
            # Extract key information from prompt
            lines = prompt.split('\n')
            candidate_section = False
            candidate_data = {}
            
            for line in lines:
                if "CANDIDATE DETAILS" in line:
                    candidate_section = True
                    continue
                elif "SCORING CRITERIA GRID" in line:
                    break
                elif candidate_section and line.strip():
                    candidate_data[line.strip()] = True
            
            # Simple scoring logic based on keywords
            java_score = 7.0 if any('java' in str(v).lower() for v in candidate_data.values()) else 5.0
            experience_score = 8.0 if any('experience' in str(v).lower() for v in candidate_data.values()) else 6.0
            education_score = 7.5 if any('bachelor' in str(v).lower() or 'degree' in str(v).lower() for v in candidate_data.values()) else 6.0
            
            final_score = (java_score + experience_score + education_score) / 3
            
            if final_score >= 8.0:
                recommendation = "To be interviewed"
            elif final_score >= 6.0:
                recommendation = "Review further"
            else:
                recommendation = "Candidature rejected"
            
            return json.dumps({
                "parameter_scores": {
                    "Java Skills": {"score": java_score, "rating": java_score * 0.6, "weightage": 60.0},
                    "Experience": {"score": experience_score, "rating": experience_score * 0.2, "weightage": 20.0},
                    "Education": {"score": education_score, "rating": education_score * 0.2, "weightage": 20.0}
                },
                "final_score": final_score,
                "recommendation": recommendation,
                "consideration": f"Based on intelligent analysis. Java: {java_score}, Experience: {experience_score}, Education: {education_score}"
            })
            
        except Exception as e:
            logger.error(f"Intelligent scoring analysis error: {e}")
            # Return a default scoring result
            return json.dumps({
                "parameter_scores": {
                    "Technical Skills": {"score": 6.0, "rating": 3.6, "weightage": 60.0},
                    "Experience": {"score": 6.0, "rating": 1.2, "weightage": 20.0},
                    "Education": {"score": 6.0, "rating": 1.2, "weightage": 20.0}
                },
                "final_score": 6.0,
                "recommendation": "Review further",
                "consideration": "Default scoring applied due to analysis error"
            })

    def _parse_scoring_response(self, response_text):
        """Parse scoring response from LLM"""
        try:
            # Try to extract JSON from the response
            import re
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                return json.loads(json_str)
            else:
                raise Exception("No JSON found in response")
        except Exception as e:
            logger.error(f"Error parsing scoring response: {e}")
            logger.error(f"Raw response: {response_text}")
            raise Exception(f"Failed to parse scoring response: {e}")

    def calculate_resume_score(self, analysis_id, criteria_id, jd_id, resume_filename, analysis_data):
        """Calculate resume score based on criteria grid and analysis data"""
        try:
            logger.info(f"Calculating score for resume: {resume_filename}")
            
            # Get resolved JD parameters
            jd_data = self.get_job_description_requirements(jd_id)
            if not jd_data:
                logger.warning(f"No JD parameters found for JD: {jd_id}")
                return None
            
            # Extract the requirements from the returned data
            jd_parameters = jd_data.get('requirements', {})
            if not jd_parameters:
                logger.warning(f"No requirements found in JD data for JD: {jd_id}")
                return None
            
            # Get criteria grid
            criteria_grid = self.get_criteria_grid(criteria_id)
            if not criteria_grid:
                logger.warning(f"No criteria grid found for criteria: {criteria_id}")
                return None
            
            # Prepare payload for scoring
            payload = {
                'jd_parameters': jd_parameters,
                'candidate_details': analysis_data,
                'criteria_grid': criteria_grid,
                'resume_filename': resume_filename
            }
            
            # Generate scoring prompt
            prompt = self._get_scoring_prompt(
                jd_parameters, 
                analysis_data, 
                criteria_grid, 
                resume_filename
            )
            
            # Call LLM for scoring
            logger.info(f"Sending scoring request to AI")
            response_text = self._call_scoring_llm(prompt)
            
            # Parse the response
            scoring_result = self._parse_scoring_response(response_text)
            logger.info(f"AI scoring completed successfully")
            
            # Save scoring result to database
            score_id = self.save_scoring_result(
                analysis_id, criteria_id, jd_id, resume_filename, scoring_result
            )
            
            return {
                'score_id': score_id,
                'scoring_result': scoring_result
            }
                
        except Exception as e:
            logger.error(f"Error calculating resume score: {e}")
            raise Exception(f"Failed to calculate resume score: {e}")

    def save_scoring_result(self, analysis_id, criteria_id, jd_id, resume_filename, scoring_result):
        """Save scoring result to database"""
        conn = self.get_db_connection()
        if not conn:
            raise Exception("Database connection failed")

        try:
            cur = conn.cursor(cursor_factory=RealDictCursor)
            
            # Extract scoring data
            parameter_scores = scoring_result.get('parameter_scores', {})
            final_score = scoring_result.get('final_score', 0)
            recommendation = scoring_result.get('recommendation', 'Review further')
            consideration = scoring_result.get('consideration', '')
            
            # Log the extracted data
            logger.info(f"Saving scoring data: final_score={final_score}, recommendation={recommendation}")
            logger.info(f"Parameter scores: {len(parameter_scores)} parameters")
            
            insert_query = """
                INSERT INTO resume_scores (
                    analysis_id, criteria_id, jd_id, resume_filename,
                    parameter_scores, final_score, recommendation, consideration
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING score_id
            """
            
            cur.execute(insert_query, (
                analysis_id,
                criteria_id,
                jd_id,
                resume_filename,
                Json(parameter_scores),
                final_score,
                recommendation,
                consideration
            ))
            
            result = cur.fetchone()
            conn.commit()
            cur.close()
            conn.close()
            
            logger.info(f"Successfully saved scoring result with ID: {result['score_id']}")
            return result['score_id']
            
        except Exception as e:
            if conn:
                conn.rollback()
                cur.close()
                conn.close()
            logger.error(f"Error saving scoring result: {e}")
            raise Exception(f"Failed to save scoring result: {e}")

@app.route('/analyze_resumes', methods=['POST'])
def analyze_resumes():
    """Main endpoint for analyzing multiple resumes against a job description"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        jd_id = data.get('jd_id')
        criteria_id = data.get('criteria_id')  # Get criteria ID for scoring
        resume_urls = data.get('resume_urls', [])
        resume_files = data.get('resume_files', [])
        
        if not jd_id:
            return jsonify({'error': 'Job description ID is required'}), 400
        
        if not resume_urls and not resume_files:
            return jsonify({'error': 'No resume URLs or files provided'}), 400
        
        # Initialize analyzer
        analyzer = ResumeAnalyzer()
        
        # Get job description requirements
        logger.info(f"Getting requirements for JD: {jd_id}")
        jd_data = analyzer.get_job_description_requirements(jd_id)
        
        if not jd_data:
            return jsonify({'error': 'Job description requirements not found'}), 404
        
        job_requirements = jd_data['requirements']
        logger.info(f"Retrieved job requirements: {job_requirements}")
        
        # Process each resume
        results = []
        failed_resumes = []
        
        # Process URLs first (files uploaded to Supabase)
        for url in resume_urls:
            try:
                logger.info(f"Processing resume URL: {url}")
                resume_text = analyzer.extract_text_from_url(url)
                
                if resume_text:
                    # Analyze the resume
                    analysis_result = analyzer.analyze_resume(resume_text, job_requirements)
                    
                    # Structure the analysis data for frontend
                    structured_analysis = analyzer._structure_analysis_data(analysis_result)
                    
                    # Save to database
                    filename = url.split('/')[-1] if '/' in url else 'resume.pdf'
                    analysis_record = analyzer.save_resume_analysis(jd_id, filename, analysis_result, url)
                    
                    # Perform scoring if criteria is provided
                    scoring_result = None
                    if criteria_id and analysis_record:
                        try:
                            logger.info(f"Starting scoring for {filename}")
                            scoring_result = analyzer.calculate_resume_score(
                                analysis_record['analysis_id'],
                                criteria_id,
                                jd_id,
                                filename,
                                analysis_result
                            )
                            logger.info(f"Scoring completed for {filename}")
                        except Exception as e:
                            logger.error(f"Scoring failed for {filename}: {e}")
                    
                    results.append({
                        'filename': filename,
                        'url': url,
                        'analysis': structured_analysis,
                        'scoring': scoring_result
                    })
                    logger.info(f"Successfully analyzed resume: {filename}")
                else:
                    failed_resumes.append({
                        'url': url,
                        'error': 'Could not extract text from resume'
                    })
                    
            except Exception as e:
                logger.error(f"Error processing resume URL {url}: {e}")
                failed_resumes.append({
                    'url': url,
                    'error': str(e)
                })
        
        # Process local files (fallback when Supabase upload fails)
        for file_path in resume_files:
            try:
                logger.info(f"Processing resume file: {file_path}")
                resume_text = analyzer.extract_text_from_local_file(file_path)
                
                if resume_text:
                    # Analyze the resume
                    analysis_result = analyzer.analyze_resume(resume_text, job_requirements)
                    
                    # Structure the analysis data for frontend
                    structured_analysis = analyzer._structure_analysis_data(analysis_result)
                    
                    # Save to database
                    filename = os.path.basename(file_path)
                    analysis_record = analyzer.save_resume_analysis(jd_id, filename, analysis_result)
                    
                    # Perform scoring if criteria is provided
                    scoring_result = None
                    if criteria_id and analysis_record:
                        try:
                            logger.info(f"Starting scoring for {filename}")
                            scoring_result = analyzer.calculate_resume_score(
                                analysis_record['analysis_id'],
                                criteria_id,
                                jd_id,
                                filename,
                                analysis_result
                            )
                            logger.info(f"Scoring completed for {filename}")
                        except Exception as e:
                            logger.error(f"Scoring failed for {filename}: {e}")
                    
                    results.append({
                        'filename': filename,
                        'file_path': file_path,
                        'analysis': structured_analysis,
                        'scoring': scoring_result
                    })
                    logger.info(f"Successfully analyzed resume: {filename}")
                else:
                    failed_resumes.append({
                        'file_path': file_path,
                        'error': 'Could not extract text from resume'
                    })
                    
            except Exception as e:
                logger.error(f"Error processing resume file {file_path}: {e}")
                failed_resumes.append({
                    'file_path': file_path,
                    'error': str(e)
                })
        
        logger.info(f"Analysis completed. Success: {len(results)}, Failed: {len(failed_resumes)}")
        logger.info(f"Input - URLs: {len(resume_urls)}, Files: {len(resume_files)}")
        logger.info(f"Output - Results: {len(results)}, Failed: {len(failed_resumes)}")
        logger.info(f"Resume URLs: {resume_urls}")
        logger.info(f"Resume Files: {resume_files}")
        
        return jsonify({
            'status': 'success',
            'message': f'Successfully analyzed {len(results)} resumes',
            'jd_id': jd_id,
            'successful_analyses': len(results),
            'failed_analyses': len(failed_resumes),
            'input_urls': len(resume_urls),
            'input_files': len(resume_files),
            'results': results,
            'failed_resumes': failed_resumes
        })
        
    except Exception as e:
        logger.error(f"Resume analysis endpoint error: {e}")
        return jsonify({
            'status': 'error',
            'message': str(e)
        }), 500

@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    try:
        analyzer = ResumeAnalyzer()
        return jsonify({
            'status': 'healthy', 
            'service': 'cv_analyzer',
            'ai_provider': analyzer.ai_provider
        })
    except Exception as e:
        return jsonify({
            'status': 'unhealthy',
            'service': 'cv_analyzer',
            'error': str(e)
        }), 500

@app.route('/')
def home():
    """Home endpoint"""
    try:
        analyzer = ResumeAnalyzer()
        return jsonify({
            'message': 'CV Analyzer Backend is running',
            'ai_provider': analyzer.ai_provider,
            'endpoints': {
                'analyze_resumes': '/analyze_resumes (POST)',
                'health': '/health (GET)'
            }
        })
    except Exception as e:
        return jsonify({
            'message': 'CV Analyzer Backend is running but AI setup failed',
            'error': str(e),
            'endpoints': {
                'analyze_resumes': '/analyze_resumes (POST)',
                'health': '/health (GET)'
            }
        })

if __name__ == '__main__':
    print("Starting CV Analyzer Backend on http://127.0.0.1:5002")
    try:
        analyzer = ResumeAnalyzer()
        print(f"AI Provider: {analyzer.ai_provider}")
    except Exception as e:
        print(f"AI Setup Error: {e}")
    app.run(debug=True, host='127.0.0.1', port=5002)